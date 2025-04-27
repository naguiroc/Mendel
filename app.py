
import streamlit as st
import pandas as pd
import re
from fractions import Fraction
from io import BytesIO
from docx import Document

# ------------------------------------------------------------------------------------------------
# FONCTIONS UTILES
# ------------------------------------------------------------------------------------------------

def underline(text):
    """
    Souligne chaque caractère d'une chaîne via un caractère combiné.
    Exemple : 'AB' -> 'A̲B̲'.
    """
    return "".join([c + "\u0332" for c in text])


def pheno(alleles, dom):
    """
    Détermine le phénotype d'une paire d'allèles selon le mode de dominance.
    - alleles : liste de deux allèles, ex. ['A', 'a']
    - dom : 'dominance complète' ou 'codominance'

    Retourne l'allèle exprimé ou 'A/a' en cas de codominance.
    """
    if dom == "dominance complète":
        return alleles[0]
    return alleles[0] if alleles[0] == alleles[1] else f"{alleles[0]}/{alleles[1]}"


def gen_gametes(gen):
    """
    Génère les gamètes et leurs fréquences pour un génotype à deux loci.
    - gen : chaîne de 4 caractères, ex. 'AaBb'
    - link, rec_rate : variables globales Streamlit

    Retourne un dict {gamete: fréquence}.
    """
    a, A, b, B = gen[0], gen[1], gen[2], gen[3]
    freqs = {}

    if link == "liés":
        if a == A and b == B:
            freqs[a + b] = 1.0
        else:
            parentaux = [a + b, A + B]
            recombines = [a + B, A + b]
            freqs[parentaux[0]] = (1 - rec_rate) / 2
            freqs[parentaux[1]] = (1 - rec_rate) / 2
            freqs[recombines[0]] = rec_rate / 2
            freqs[recombines[1]] = rec_rate / 2
    else:
        # Pour gènes indépendants, on cumule fréquences pour gérer homozygotes
        for X in [a, A]:
            for Y in [b, B]:
                gam = X + Y
                freqs[gam] = freqs.get(gam, 0) + 1/4
    return freqs

# ------------------------------------------------------------------------------------------------
# CONFIGURATION STREAMLIT
# ------------------------------------------------------------------------------------------------

st.set_page_config(page_title="Échiquier de croisement", layout="centered")
st.title("🧬 Interprétation chromosomique :")
st.markdown("Saisissez les génotypes parentaux, modes de transmission et taux de recombinaison si applicable.")

if 'generated' not in st.session_state:
    st.session_state.generated = False

def reset_state():
    st.session_state.generated = False

# ------------------------------------------------------------------------------------------------
# WIDGETS DE SAISIE
# ------------------------------------------------------------------------------------------------

p1 = st.text_input("Génotype du Parent 1 (ex: AaBb)", "AaBb", on_change=reset_state)
p2 = st.text_input("Génotype du Parent 2 (ex: AaBb)", "AaBb", on_change=reset_state)

domA = st.selectbox("Dominance du gène A", ["dominance complète", "codominance"], on_change=reset_state)
domB = st.selectbox("Dominance du gène B", ["dominance complète", "codominance"], on_change=reset_state)

link = st.radio("Les deux gènes sont :", ["indépendants", "liés"], horizontal=True, on_change=reset_state)
rec_rate = 0.0
if link == "liés":
    rec_rate = st.slider("Taux de recombinaison (r)", 0.0, 0.5, 0.1, 0.01, on_change=reset_state)

geno_pattern = re.compile(r'^[AaBb]{4}$')

def generate_board():
    if not (geno_pattern.match(p1) and geno_pattern.match(p2)):
        st.error("Format des génotypes invalide : utilisez AaBb, AABB, aaBb, etc.")
    else:
        st.session_state.generated = True

st.button("Générer l’échiquier", on_click=generate_board)

# ------------------------------------------------------------------------------------------------
# GÉNÉRATION ET AFFICHAGE
# ------------------------------------------------------------------------------------------------

if st.session_state.generated:
    # Création document Word
    doc = Document()
    doc.add_heading("Interprétation chromosomique du croisement", level=1)

    # Séparer allèles
    A1, B1 = sorted([p1[0], p1[1]]), sorted([p1[2], p1[3]])
    A2, B2 = sorted([p2[0], p2[1]]), sorted([p2[2], p2[3]])
    phA1, phB1 = pheno(A1, domA), pheno(B1, domB)
    phA2, phB2 = pheno(A2, domA), pheno(B2, domB)

    # Générer gamètes
    g1 = gen_gametes(p1)
    g2 = gen_gametes(p2)

    # Cas homozygotes stricts (AABB x aabb) : Affichage complet des parents, phénotype, génotype, gamètes puis descendance
    if len(g1) == 1 and len(g2) == 1:
        # Affichage des parents
        st.markdown("**Les parents**")
        if link == "indépendants":
            geno1 = f"{A1[0]}//{A1[1]} {B1[0]}//{B1[1]}"
            geno2 = f"{A2[0]}//{A2[1]} {B2[0]}//{B2[1]}"
        else:
            geno1 = f"{p1[0]}{p1[1]}//{p1[2]}{p1[3]}"
            geno2 = f"{p2[0]}{p2[1]}//{p2[2]}{p2[3]}"
        indent = 50
        # Phénotype et génotype décalés
        st.markdown(
            f"<div style='padding-left:{indent}px'><strong>Phénotype :</strong> [{phA1},{phB1}] × [{phA2},{phB2}]</div>",
            unsafe_allow_html=True
        )
        st.markdown(
            f"<div style='padding-left:{indent}px'><strong>Génotype :</strong> {geno1} × {geno2}</div>",
            unsafe_allow_html=True
        )
        # Génération et affichage des gamètes
        def fmt_g(g, f): return f"{f:.2f} {underline(g)}"
        gam_p1 = ", ".join(fmt_g(k, v) for k, v in g1.items())
        gam_p2 = ", ".join(fmt_g(k, v) for k, v in g2.items())
        st.markdown(f"**Les gamètes :** {gam_p1}   x   {gam_p2}")
        # Descendance à 100%
        # Construction du génotype enfant
        gam1, _ = next(iter(g1.items()))
        gam2, _ = next(iter(g2.items()))
        AA = sorted([gam1[0], gam2[0]])
        BB = sorted([gam1[1], gam2[1]])
        geno_desc = f"{AA[0]}//{AA[1]} {BB[0]}//{BB[1]}"
        phA_desc = pheno(AA, domA)
        phB_desc = pheno(BB, domB)
        st.markdown(f"**Descendance :** {geno_desc} [{phA_desc},{phB_desc}] 100%")
        # Écriture dans Word
        doc.add_heading("Les parents", level=2)
        doc.add_paragraph(f"Phénotype : [{phA1},{phB1}] × [{phA2},{phB2}]")
        doc.add_paragraph(f"Génotype : {geno1} × {geno2}")
        doc.add_heading("Gamètes", level=2)
        doc.add_paragraph(f"{gam_p1}   x   {gam_p2}")
        doc.add_heading("Descendance", level=2)
        doc.add_paragraph(f"{geno_desc} [{phA_desc},{phB_desc}] 100%")
    else:
        # Afficher parents
        st.markdown("**Les parents**")
        if link == "indépendants":
            geno1 = f"{A1[0]}//{A1[1]} {B1[0]}//{B1[1]}"
            geno2 = f"{A2[0]}//{A2[1]} {B2[0]}//{B2[1]}"
        else:
            geno1 = f"{p1[0]}{p1[1]}//{p1[2]}{p1[3]}"
            geno2 = f"{p2[0]}{p2[1]}//{p2[2]}{p2[3]}"
        indent = 50
        st.markdown(f"<div style='padding-left:{indent}px'><strong>Phénotype :</strong> [{phA1},{phB1}] × [{phA2},{phB2}]</div>", unsafe_allow_html=True)
        st.markdown(f"<div style='padding-left:{indent}px'><strong>Génotype :</strong> {geno1} × {geno2}</div>", unsafe_allow_html=True)
        doc.add_paragraph(f"Phénotype : [{phA1},{phB1}] × [{phA2},{phB2}]")
        doc.add_paragraph(f"Génotype : {geno1} × {geno2}")

        # Afficher gamètes
        def fmt_g(g, f): return f"{f:.2f} {underline(g)}"
        gam_p1 = ", ".join(fmt_g(k, v) for k, v in g1.items())
        gam_p2 = ", ".join(fmt_g(k, v) for k, v in g2.items())
        st.markdown(f"**Les gamètes :** {gam_p1}   x   {gam_p2}")
        doc.add_heading("Gamètes", level=2)
        doc.add_paragraph(f"{gam_p1}   x   {gam_p2}")

        # Génération de l'échiquier
        st.markdown("## L'échiquier de croisement")
        # Détection Punnett horizontal vs 4x4
        horizontal = (len(g1) == 1 and len(g2) == 4) or (len(g2) == 1 and len(g1) == 4)
        # Préparer DataFrame
        if horizontal:
            outer, inner = (g1, g2) if len(g1) == 1 else (g2, g1)
            index = [fmt_g(k, v) for k, v in outer.items()]
            columns = [fmt_g(k, v) for k, v in inner.items()]
            rows = []
            for o_g, o_f in outer.items():
                row = []
                for i_g, i_f in inner.items():
                    AA = sorted([o_g[0], i_g[0]])
                    BB = sorted([o_g[1], i_g[1]])
                    geno = f"{AA[0]}//{AA[1]} {BB[0]}//{BB[1]}" if link=="indépendants" else f"{min(o_g,i_g)}//{max(o_g,i_g)}"
                    phA = pheno(AA, domA)
                    phB = pheno(BB, domB)
                    frac = str(Fraction(o_f*i_f).limit_denominator())
                    # Deux lignes en cellule
                    cell = f"{geno}\n({frac}) [{phA},{phB}]"
                    row.append(cell)
                rows.append(row)
            df = pd.DataFrame(rows, index=index, columns=columns)
        else:
            rows = []
            for a_g, a_f in g1.items():
                row = []
                for b_g, b_f in g2.items():
                    AA = sorted([a_g[0], b_g[0]])
                    BB = sorted([a_g[1], b_g[1]])
                    geno = f"{AA[0]}//{AA[1]} {BB[0]}//{BB[1]}" if link=="indépendants" else f"{min(a_g,b_g)}//{max(a_g,b_g)}"
                    phA = pheno(AA, domA)
                    phB = pheno(BB, domB)
                    frac = str(Fraction(a_f*b_f).limit_denominator())
                    cell = f"{geno}\n({frac}) [{phA},{phB}]"
                    row.append(cell)
                rows.append(row)
            df = pd.DataFrame(
                rows,
                index=[fmt_g(k,v) for k,v in g1.items()],
                columns=[fmt_g(k,v) for k,v in g2.items()]
            )

        # Affichage Streamlit
        st.table(df.style.set_table_styles([
            {'selector':'td,th','props':[('border','1px solid black')]}
        ]))

        # Export Word
        doc.add_heading("Échiquier de croisement", level=2)
        ncol = len(df.columns)+1
        tbl = doc.add_table(rows=1, cols=ncol)
        hdr = tbl.rows[0].cells
        # angle en-tête
        hdr[0].text = "↓ γ♂   ↑ γ♀ →"
        for i, c in enumerate(df.columns,1): hdr[i].text = c
        for ridx, row in enumerate(df.itertuples(),1):
            cells = tbl.add_row().cells
            cells[0].text = df.index[ridx-1]
            for j, val in enumerate(row[1:],1): cells[j].text = val

    # Bouton download
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    st.download_button("📄 Télécharger le rapport (DOCX)", buf, "croisement.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")